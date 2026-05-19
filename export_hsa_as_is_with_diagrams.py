from __future__ import annotations

import math
import re
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Pt, RGBColor


SRC = Path("hsa-education-as-is-operations-analysis.md")
OUT = Path("hsa-education-as-is-operations-analysis-landscape-diagrams.docx")
DIAGRAM_DIR = Path("hsa-as-is-models-png")

W = 2600
BG = (250, 252, 254)
INK = (34, 39, 46)
MUTED = (94, 103, 112)
NAVY = (31, 78, 121)
BLUE = (67, 125, 191)
TEAL = (32, 151, 144)
GREEN = (76, 160, 102)
GOLD = (218, 151, 48)
ORANGE = (229, 116, 69)
RED = (198, 72, 72)
PURPLE = (116, 92, 165)
LINE = (196, 207, 219)
WHITE = (255, 255, 255)


DIAGRAM_TITLES = [
    "Sơ đồ tổ chức cấp cao",
    "Cơ cấu Sale theo miền",
    "Bảng tổng kết nhân sự",
    "Bản đồ luồng dữ liệu giữa các hệ thống",
    "Luồng 1 - Marketing & Tạo lead",
    "Luồng 2 - CRM, Sale và Nurture",
    "Luồng hỗ trợ case khó trong Sale/CTV",
    "Luồng 3 - Thanh toán",
    "Luồng 4 - Onboarding sau thanh toán",
    "Luồng 5 - Học tập trên ClassIn/Zalo",
    "Luồng 6 - Chăm sóc học viên",
    "Luồng 7 - Quản lý giảng viên và thù lao",
    "Luồng 8 - Quản lý CTV Sale và hoa hồng",
    "Luồng 9 - Đối soát kế toán",
    "Phụ thuộc chéo giữa các luồng vận hành",
]


def font(size: int, bold: bool = False, italic: bool = False) -> ImageFont.FreeTypeFont:
    candidates = []
    if bold:
        candidates += [
            Path("C:/Windows/Fonts/arialbd.ttf"),
            Path("C:/Windows/Fonts/segoeuib.ttf"),
        ]
    elif italic:
        candidates += [
            Path("C:/Windows/Fonts/ariali.ttf"),
            Path("C:/Windows/Fonts/segoeuii.ttf"),
        ]
    candidates += [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/segoeui.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def canvas(height: int) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (W, height), BG)
    draw = ImageDraw.Draw(img)
    return img, draw


def text_box(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont) -> tuple[int, int]:
    if not text:
        return (0, 0)
    box = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=8)
    return (box[2] - box[0], box[3] - box[1])


def wrap_line(draw: ImageDraw.ImageDraw, line: str, fnt: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    words = line.split()
    if not words:
        return [""]
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if text_box(draw, candidate, fnt)[0] <= max_width:
            current = candidate
            continue
        if current:
            lines.append(current)
            current = word
        else:
            chunk = ""
            for ch in word:
                candidate_chunk = f"{chunk}{ch}"
                if text_box(draw, candidate_chunk, fnt)[0] <= max_width:
                    chunk = candidate_chunk
                else:
                    if chunk:
                        lines.append(chunk)
                    chunk = ch
            current = chunk
    if current:
        lines.append(current)
    return lines


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    wrapped: list[str] = []
    for line in text.splitlines():
        wrapped.extend(wrap_line(draw, line.strip(), fnt, max_width))
    return wrapped or [""]


def dashed_line(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    fill: tuple[int, int, int],
    width: int = 5,
    dash: int = 24,
    gap: int = 14,
) -> None:
    x1, y1 = start
    x2, y2 = end
    length = math.hypot(x2 - x1, y2 - y1)
    if length == 0:
        return
    dx = (x2 - x1) / length
    dy = (y2 - y1) / length
    distance = 0
    while distance < length:
        segment = min(dash, length - distance)
        sx = x1 + dx * distance
        sy = y1 + dy * distance
        ex = x1 + dx * (distance + segment)
        ey = y1 + dy * (distance + segment)
        draw.line((sx, sy, ex, ey), fill=fill, width=width)
        distance += dash + gap


def arrow_head(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    fill: tuple[int, int, int],
    size: int = 22,
) -> None:
    x1, y1 = start
    x2, y2 = end
    angle = math.atan2(y2 - y1, x2 - x1)
    p1 = (x2, y2)
    p2 = (x2 - size * math.cos(angle - math.pi / 6), y2 - size * math.sin(angle - math.pi / 6))
    p3 = (x2 - size * math.cos(angle + math.pi / 6), y2 - size * math.sin(angle + math.pi / 6))
    draw.polygon([p1, p2, p3], fill=fill)


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    fill: tuple[int, int, int] = NAVY,
    width: int = 6,
    label: str | None = None,
    label_offset: tuple[int, int] = (0, 0),
    dashed: bool = False,
) -> None:
    if dashed:
        dashed_line(draw, start, end, fill, width=width)
    else:
        draw.line((*start, *end), fill=fill, width=width)
    arrow_head(draw, start, end, fill=fill, size=26)
    if label:
        mx = (start[0] + end[0]) // 2 + label_offset[0]
        my = (start[1] + end[1]) // 2 + label_offset[1]
        draw_badge(draw, (mx, my), label, fill=(255, 250, 235), outline=(236, 181, 79), text_color=INK)


def poly_arrow(
    draw: ImageDraw.ImageDraw,
    points: list[tuple[int, int]],
    fill: tuple[int, int, int] = NAVY,
    width: int = 6,
    dashed: bool = False,
    label: str | None = None,
    label_at: tuple[int, int] | None = None,
) -> None:
    for a, b in zip(points, points[1:]):
        if dashed:
            dashed_line(draw, a, b, fill, width=width)
        else:
            draw.line((*a, *b), fill=fill, width=width)
    arrow_head(draw, points[-2], points[-1], fill=fill, size=26)
    if label and label_at:
        draw_badge(draw, label_at, label, fill=(255, 244, 232), outline=(232, 162, 112), text_color=INK)


def draw_badge(
    draw: ImageDraw.ImageDraw,
    center: tuple[int, int],
    text: str,
    fill: tuple[int, int, int],
    outline: tuple[int, int, int],
    text_color: tuple[int, int, int] = INK,
    size: int = 31,
) -> None:
    fnt = font(size, bold=True)
    pad_x, pad_y = 20, 10
    tw, th = text_box(draw, text, fnt)
    x, y = center
    rect = (x - tw // 2 - pad_x, y - th // 2 - pad_y, x + tw // 2 + pad_x, y + th // 2 + pad_y)
    draw.rounded_rectangle(rect, radius=24, fill=fill, outline=outline, width=3)
    draw.text((x - tw // 2, y - th // 2 - 2), text, font=fnt, fill=text_color)


def draw_title(draw: ImageDraw.ImageDraw, title: str, subtitle: str | None = None) -> None:
    draw.rectangle((0, 0, W, 22), fill=NAVY)
    title_font = font(58, bold=True)
    draw.text((100, 58), title, font=title_font, fill=NAVY)
    if subtitle:
        draw.text((104, 128), subtitle, font=font(32), fill=MUTED)
    draw.line((100, 178, W - 100, 178), fill=LINE, width=4)


def draw_footer(draw: ImageDraw.ImageDraw, height: int, number: int) -> None:
    draw.line((100, height - 70, W - 100, height - 70), fill=(222, 228, 236), width=3)
    draw.text((100, height - 48), f"Hình {number:02d} | HSA Education As-Is Operations Analysis Q2/2026", font=font(24), fill=MUTED)


def draw_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    text: str,
    fill: tuple[int, int, int] = WHITE,
    outline: tuple[int, int, int] = NAVY,
    text_color: tuple[int, int, int] = INK,
    size: int = 38,
    bold: bool = False,
    radius: int = 28,
    align: str = "center",
    stroke: int = 4,
    padding: int = 28,
) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=stroke)
    fnt_size = size
    while fnt_size >= 21:
        fnt = font(fnt_size, bold=bold)
        lines = wrap_text(draw, text, fnt, x2 - x1 - padding * 2)
        line_h = text_box(draw, "Ag", fnt)[1] + 8
        total_h = line_h * len(lines)
        max_w = max(text_box(draw, line, fnt)[0] for line in lines)
        if total_h <= y2 - y1 - padding * 2 and max_w <= x2 - x1 - padding * 2:
            break
        fnt_size -= 2
    top = y1 + (y2 - y1 - total_h) // 2
    for idx, line in enumerate(lines):
        tw, th = text_box(draw, line, fnt)
        if align == "left":
            tx = x1 + padding
        else:
            tx = x1 + (x2 - x1 - tw) // 2
        draw.text((tx, top + idx * line_h), line, font=fnt, fill=text_color)


def render_org(number: int, path: Path) -> None:
    img, draw = canvas(1450)
    draw_title(draw, "Sơ đồ tổ chức cấp cao", "Cấu trúc lãnh đạo và các khối chức năng hiện tại")
    draw_box(draw, (865, 235, 1735, 390), "HĐQT / Đồng sáng lập\nHoa · Thầy Khương", fill=(235, 243, 252), outline=NAVY, size=40, bold=True)
    draw_box(draw, (390, 555, 1110, 735), "GĐ Vận hành Bắc\n3 kỳ thi: HSA · BCA · BQP", fill=(236, 248, 246), outline=TEAL, size=38, bold=True)
    draw_box(draw, (1490, 555, 2210, 735), "GĐ Vận hành Nam\n1 kỳ thi: ĐGNL HCM", fill=(255, 245, 232), outline=GOLD, size=38, bold=True)
    arrow(draw, (1300, 390), (1300, 500), fill=NAVY)
    draw.line((750, 500, 1850, 500), fill=NAVY, width=6)
    arrow(draw, (750, 500), (750, 555), fill=NAVY)
    arrow(draw, (1850, 500), (1850, 555), fill=NAVY)
    draw.line((750, 735, 750, 855), fill=LINE, width=5)
    draw.line((1850, 735, 1850, 855), fill=LINE, width=5)
    draw.line((475, 855, 2125, 855), fill=LINE, width=5)
    functions = [
        ("Kế toán\n(chung)", 285, BLUE),
        ("Sale", 675, TEAL),
        ("Học vụ\n& QLL", 1065, GREEN),
        ("Truyền\nthông", 1455, PURPLE),
        ("Hành chính\n& NS", 1845, ORANGE),
    ]
    for text, x, color in functions:
        draw_box(draw, (x, 965, x + 340, 1125), text, fill=WHITE, outline=color, size=34, bold=True)
        arrow(draw, (x + 170, 855), (x + 170, 965), fill=LINE, width=5)
    draw_badge(draw, (1300, 1230), "2 miền vận hành, nhiều chức năng dùng chung", fill=(244, 247, 251), outline=LINE, size=34)
    draw_badge(draw, (1850, 1320), "HCM: offline 12 | Sale 20-25 | Marketing 20 | Vận hành lớp 10 | GV chính 15", fill=(255, 245, 232), outline=GOLD, size=27)
    draw_footer(draw, 1450, number)
    img.save(path, dpi=(180, 180))


def render_sales_structure(number: int, path: Path) -> None:
    img, draw = canvas(1280)
    draw_title(draw, "Cơ cấu Sale theo miền", "Hai miền là 2 cụm song song; các box dưới mỗi miền là nhánh, không phải chuỗi bước")
    draw_box(draw, (170, 265, 1210, 415), "MIỀN BẮC (HÀ NỘI)", fill=(235, 243, 252), outline=NAVY, size=44, bold=True)
    draw_box(draw, (1390, 265, 2430, 415), "MIỀN NAM (HỒ CHÍ MINH)", fill=(255, 245, 232), outline=GOLD, size=44, bold=True)

    draw_box(draw, (170, 600, 660, 760), "Phòng Sale Offline\n1 Trưởng phòng + 11 sale\n= 12 người", fill=WHITE, outline=TEAL, size=30, bold=True)
    draw_box(draw, (720, 600, 1210, 760), "Đội Sale Online\n1 quản lý + ~100 CTV", fill=WHITE, outline=PURPLE, size=31, bold=True)
    draw.line((690, 415, 690, 515), fill=LINE, width=5)
    draw.line((415, 515, 965, 515), fill=LINE, width=5)
    arrow(draw, (415, 515), (415, 600), fill=LINE, width=5)
    arrow(draw, (965, 515), (965, 600), fill=LINE, width=5)

    draw_box(draw, (1505, 600, 2315, 760), "Sale HCM\n20-25 người", fill=WHITE, outline=ORANGE, size=36, bold=True)
    arrow(draw, (1910, 415), (1910, 600), fill=LINE, width=5)
    draw_box(draw, (1505, 820, 2315, 940), "Cần tách rõ fulltime / CTV\nđể phân quyền CRM và tính hoa hồng", fill=(255, 246, 246), outline=RED, size=28, bold=True)

    draw_box(draw, (390, 1030, 2210, 1148), "TỔNG MẠNG LƯỚI SALE: ~132-137 người\nHN ~112 người + HCM 20-25 người", fill=(244, 247, 251), outline=NAVY, size=39, bold=True)
    draw_footer(draw, 1280, number)
    img.save(path, dpi=(180, 180))


def panel(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, color: tuple[int, int, int]) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=30, fill=WHITE, outline=color, width=4)
    draw.rounded_rectangle((x1, y1, x2, y1 + 88), radius=30, fill=color, outline=color, width=4)
    draw.rectangle((x1, y1 + 45, x2, y1 + 88), fill=color)
    tw, th = text_box(draw, title, font(34, bold=True))
    draw.text((x1 + (x2 - x1 - tw) // 2, y1 + 25), title, font=font(34, bold=True), fill=WHITE)


def draw_row(draw: ImageDraw.ImageDraw, x: int, y: int, right_x: int, label: str, value: str, color: tuple[int, int, int]) -> None:
    label_font = font(27)
    value_font = font(27, bold=True)
    draw.text((x, y), label, font=label_font, fill=INK)
    tw, th = text_box(draw, value, value_font)
    badge_w = max(76, tw + 34)
    badge_h = 50
    bx2 = right_x
    bx1 = bx2 - badge_w
    draw.rounded_rectangle((bx1, y - 8, bx2, y - 8 + badge_h), radius=17, fill=(246, 248, 250), outline=(226, 231, 238), width=2)
    draw.text((bx1 + (badge_w - tw) // 2, y + (badge_h - th) // 2 - 6), value, font=value_font, fill=color)


def render_headcount(number: int, path: Path) -> None:
    img, draw = canvas(1280)
    draw_title(draw, "Bảng tổng kết nhân sự", "Quy mô nhân lực đang vận hành trong Q2/2026")
    draw_box(draw, (740, 225, 1860, 345), "TỔNG TOÀN TỔ CHỨC: >300 người", fill=(235, 243, 252), outline=NAVY, size=43, bold=True)
    panels = [
        ((100, 430, 640, 810), "FULLTIME / OFFLINE", NAVY, [
            ("Hà Nội", "50"),
            ("Hồ Chí Minh", "12"),
            ("Tổng offline", "62"),
        ]),
        ((700, 430, 1240, 1095), "SALE / MARKETING", PURPLE, [
            ("Sale offline HN", "12"),
            ("CTV Sale HN", "~100"),
            ("Sale HCM", "20-25"),
            ("Marketing HCM", "20"),
            ("Marketing HN", "~15"),
            ("Sale network", "~132-137"),
        ]),
        ((1310, 430, 1850, 875), "VẬN HÀNH LỚP", TEAL, [
            ("QLL hiện hữu", "8"),
            ("Vận hành HCM", "10"),
            ("Tổng vận hành", "~18"),
            ("HS/người", "~1.110"),
        ]),
        ((1920, 430, 2460, 810), "GIÁO VIÊN", GREEN, [
            ("GV online", "~70"),
            ("GV chính HCM", "15"),
            ("GV đề", "ngoài scope"),
        ]),
    ]
    for xy, title, color, rows in panels:
        panel(draw, xy, title, color)
        y = xy[1] + 132
        right_x = xy[2] - 34
        for label, value in rows:
            draw_row(draw, xy[0] + 50, y, right_x, label, value, color)
            y += 66
        if title.startswith("SALE"):
            draw_badge(draw, ((xy[0] + xy[2]) // 2, xy[3] - 54), "cần tách fulltime / CTV", fill=(250, 246, 255), outline=(204, 188, 233), size=23)
    draw_badge(draw, (1300, 1160), "Tổng là ước tính vì HCM có nhóm pha trộn fulltime và CTV, cần chuẩn hóa dữ liệu để tránh đếm trùng.", fill=(244, 247, 251), outline=LINE, size=28)
    draw_footer(draw, 1280, number)
    img.save(path, dpi=(180, 180))


def render_system_map(number: int, path: Path) -> None:
    img, draw = canvas(1600)
    draw_title(draw, "Bản đồ luồng dữ liệu giữa các hệ thống", "Layout lane rõ ràng: tự động, thủ công và điểm chưa có data pipeline")
    nodes = {
        "ez": ((120, 315, 620, 465), "EZSale\n(CRM)", BLUE),
        "web": ((1050, 315, 1550, 465), "Web portal\n(form/database)", TEAL),
        "sepay": ((1980, 315, 2480, 465), "SePay\n(thanh toán)", GREEN),
        "sheet": ((120, 745, 620, 895), "Google Sheet\n(master dữ liệu)", PURPLE),
        "zalo": ((1050, 745, 1550, 895), "Zalo\n(OA + nhóm)", ORANGE),
        "classin": ((1980, 745, 2480, 895), "ClassIn\n(lớp học live)", NAVY),
        "accounting": ((120, 1165, 620, 1315), "Kế toán\nđối soát", RED),
        "report": ((1050, 1165, 1550, 1315), "Báo cáo cuối tháng\nSheet thủ công", GOLD),
        "missing": ((1980, 1165, 2480, 1315), "Chưa có\nClassIn data pipeline", RED),
    }
    for xy, text, color in nodes.values():
        draw_box(draw, xy, text, fill=WHITE, outline=color, size=34, bold=True)
    arrow(draw, (1050, 390), (620, 390), fill=ORANGE, label="nhập tay / chưa chắc", label_offset=(0, -62), dashed=True)
    arrow(draw, (1550, 390), (1980, 390), fill=GREEN, label="webhook ổn định", label_offset=(0, -62))
    arrow(draw, (370, 465), (370, 745), fill=ORANGE, label="nhập tay", label_offset=(105, 0), dashed=True)
    arrow(draw, (1300, 465), (1300, 745), fill=ORANGE, label="gửi tay", label_offset=(105, 0), dashed=True)
    arrow(draw, (2230, 465), (2230, 745), fill=GREEN, label="payment_success", label_offset=(135, 0))
    arrow(draw, (620, 820), (1050, 820), fill=ORANGE, label="copy-paste", label_offset=(0, -62), dashed=True)
    arrow(draw, (1550, 820), (1980, 820), fill=RED, label="không kết nối sâu", label_offset=(0, -62), dashed=True)
    arrow(draw, (370, 895), (370, 1165), fill=ORANGE, label="đối soát thủ công", label_offset=(150, 0), dashed=True)
    arrow(draw, (620, 1240), (1050, 1240), fill=ORANGE, label="Sheet thủ công", label_offset=(0, -62), dashed=True)
    arrow(draw, (2230, 895), (2230, 1165), fill=RED, label="chưa sync", label_offset=(125, 0), dashed=True)
    draw_badge(draw, (820, 245), "Thủ công / rời rạc", fill=(255, 244, 232), outline=(232, 162, 112), text_color=(150, 74, 35), size=29)
    draw_badge(draw, (2080, 245), "Tự động ổn định", fill=(238, 249, 241), outline=(151, 204, 167), text_color=(31, 111, 59), size=29)
    draw_footer(draw, 1600, number)
    img.save(path, dpi=(180, 180))


def render_vertical_flow(
    number: int,
    path: Path,
    title: str,
    steps: list[str],
    notes: list[str | None] | None = None,
    highlights: dict[int, tuple[int, int, int]] | None = None,
    height: int = 1600,
) -> None:
    img, draw = canvas(height)
    draw_title(draw, title, "Luồng As-Is đang diễn ra trong vận hành")
    notes = notes or [None] * (len(steps) - 1)
    highlights = highlights or {}
    top = 240
    bottom = 110
    gap = 70
    box_h = max(105, min(175, (height - top - bottom - gap * (len(steps) - 1)) // len(steps)))
    flow_h = box_h * len(steps) + gap * (len(steps) - 1)
    y = top + max(0, (height - top - bottom - flow_h) // 2)
    x1, x2 = 430, 2170
    centers: list[tuple[int, int]] = []
    for idx, step in enumerate(steps):
        color = highlights.get(idx, TEAL if idx == 0 else NAVY)
        fill = (255, 248, 245) if color == RED else WHITE
        draw_box(draw, (x1, y, x2, y + box_h), step, fill=fill, outline=color, size=35, bold=True if idx == 0 or color == RED else False)
        centers.append(((x1 + x2) // 2, y + box_h // 2))
        if idx < len(steps) - 1:
            arrow(draw, ((x1 + x2) // 2, y + box_h), ((x1 + x2) // 2, y + box_h + gap - 12), fill=highlights.get(idx + 1, LINE), width=5, label=notes[idx], label_offset=(260, 0) if notes[idx] else (0, 0), dashed=bool(notes[idx] and "thủ công" in notes[idx].lower()))
        y += box_h + gap
    draw_footer(draw, height, number)
    img.save(path, dpi=(180, 180))


def render_crm_nurture(number: int, path: Path) -> None:
    img, draw = canvas(1550)
    draw_title(draw, "Luồng 2 - CRM, Sale và Nurture", "Phân nhánh Hot / Warm / Cold và cách chốt đơn qua Zalo cá nhân")
    draw_box(draw, (730, 250, 1870, 390), "Lead trong EZSale\nHot / Warm / Cold", fill=WHITE, outline=NAVY, size=38, bold=True)
    draw_box(draw, (240, 590, 1120, 780), "HOT\nSale gọi điện trực tiếp, chốt đơn", fill=(255, 246, 246), outline=RED, size=36, bold=True)
    draw_box(draw, (1480, 590, 2360, 820), "WARM / COLD\nSale nhắc lại thủ công\nZalo cá nhân, gọi điện, nhóm Zalo CTV để nurture", fill=(255, 249, 237), outline=GOLD, size=32, bold=True)
    arrow(draw, (1100, 390), (680, 590), fill=RED)
    arrow(draw, (1500, 390), (1920, 590), fill=GOLD, label="thủ công", label_offset=(130, -30), dashed=True)
    draw_box(draw, (730, 1010, 1870, 1165), "Khi gần chốt\nSale gửi link giỏ hàng web qua Zalo cá nhân", fill=WHITE, outline=TEAL, size=34, bold=True)
    arrow(draw, (680, 780), (1100, 1010), fill=LINE)
    arrow(draw, (1920, 820), (1500, 1010), fill=LINE)
    draw_box(draw, (730, 1280, 1870, 1400), "Học sinh tự vào web, chọn khóa, thanh toán", fill=(237, 249, 241), outline=GREEN, size=36, bold=True)
    arrow(draw, (1300, 1165), (1300, 1280), fill=GREEN)
    draw_footer(draw, 1550, number)
    img.save(path, dpi=(180, 180))


def render_accounting(number: int, path: Path) -> None:
    img, draw = canvas(1320)
    draw_title(draw, "Luồng 9 - Đối soát kế toán", "Tách rõ việc hằng ngày và cụm việc dồn vào cuối tháng")
    draw_box(draw, (150, 270, 1180, 430), "HẰNG NGÀY\nKế toán thu đối soát SePay log vs đơn hàng\n~2 giờ/ngày", fill=(235, 243, 252), outline=NAVY, size=34, bold=True)
    draw_box(draw, (1420, 270, 2450, 430), "CUỐI THÁNG\nKế toán chi + kế toán tổng hợp xử lý batch", fill=(255, 245, 232), outline=GOLD, size=34, bold=True)

    draw_box(draw, (150, 570, 1180, 720), "Đầu vào hằng ngày\nSePay log · Đơn hàng web", fill=WHITE, outline=GREEN, size=32, bold=True)
    arrow(draw, (665, 430), (665, 570), fill=LINE, width=5)
    draw_box(draw, (150, 835, 1180, 1015), "Kết quả hiện tại\nĐối soát thủ công; chưa auto-match thành báo cáo kế toán realtime", fill=WHITE, outline=RED, size=31, bold=True)
    arrow(draw, (665, 720), (665, 835), fill=ORANGE, width=5, label="thủ công", label_offset=(135, 0), dashed=True)

    monthly = [
        ("Thù lao GV", "~1 ngày", TEAL),
        ("Hoa hồng Sale/CTV", "~2 ngày", PURPLE),
        ("Báo cáo tổng hợp", "~vài ngày", ORANGE),
    ]
    y = 560
    for title, value, color in monthly:
        draw_box(draw, (1420, y, 2450, y + 135), f"{title}\n{value}", fill=WHITE, outline=color, size=33, bold=True)
        y += 180
    draw_box(draw, (360, 1120, 2240, 1215), "Rủi ro chính: báo cáo phụ thuộc thao tác thủ công và đối chiếu nhiều nguồn.", fill=(255, 246, 246), outline=RED, size=32, bold=True)
    draw_footer(draw, 1320, number)
    img.save(path, dpi=(180, 180))


def render_dependencies(number: int, path: Path) -> None:
    img, draw = canvas(1650)
    draw_title(draw, "Phụ thuộc chéo giữa các luồng vận hành", "Ma trận phụ thuộc đúng theo từng dòng, tránh fan-in và mũi tên chồng chéo")
    header_font = font(31, bold=True)
    draw.text((170, 240), "Nguồn dữ liệu / điều kiện", font=header_font, fill=NAVY)
    draw.text((960, 240), "Luồng phụ thuộc", font=header_font, fill=NAVY)
    draw.text((1710, 240), "Trạng thái hiện tại", font=header_font, fill=NAVY)
    draw.line((120, 295, 2480, 295), fill=LINE, width=4)

    rows = [
        ("Luồng 3\nSePay webhook", "Luồng 4\nOnboarding", "payment_success\nlà điều kiện bắt đầu", GREEN, ORANGE),
        ("Luồng 4\nHS đã onboard", "Luồng 5\nHọc tập", "Sau onboarding\nmới vào lớp", ORANGE, NAVY),
        ("Luồng 5\nClassIn data", "Luồng 6\nChăm sóc", "Chưa sync tự động", NAVY, PURPLE),
        ("Luồng 5\nGiờ dạy", "Luồng 7\nThù lao GV", "Tính thủ công", NAVY, BLUE),
        ("Luồng 1\nAttribution/ref", "Luồng 8\nHoa hồng Sale/CTV", "Thiếu link tracking", TEAL, RED),
        ("Luồng 3 + 7 + 8\nDoanh thu / chi phí", "Luồng 9\nKế toán", "Fan-in cuối tháng\nđối chiếu thủ công", GOLD, GOLD),
    ]
    y = 330
    for source, dependent, status, source_color, dependent_color in rows:
        draw_box(draw, (130, y, 700, y + 135), source, fill=WHITE, outline=source_color, size=30, bold=True)
        arrow(draw, (700, y + 68), (900, y + 68), fill=LINE, width=5)
        draw_box(draw, (900, y, 1470, y + 135), dependent, fill=WHITE, outline=dependent_color, size=30, bold=True)
        status_fill = (255, 246, 246) if dependent_color in (RED, BLUE, PURPLE) or "thủ công" in status.lower() or "Thiếu" in status else (255, 248, 235)
        draw_box(draw, (1680, y, 2470, y + 135), status, fill=status_fill, outline=dependent_color if dependent_color != NAVY else LINE, size=28, bold=True)
        y += 185

    draw_box(draw, (360, 1470, 2240, 1565), "Quan sát: phần lớn phụ thuộc đang được bù bằng thao tác thủ công thay vì kết nối hệ thống.", fill=(255, 246, 246), outline=RED, size=31, bold=True)
    draw_footer(draw, 1650, number)
    img.save(path, dpi=(180, 180))


def generate_diagrams() -> list[Path]:
    DIAGRAM_DIR.mkdir(exist_ok=True)
    paths = [
        DIAGRAM_DIR / "01-so-do-to-chuc-cap-cao.png",
        DIAGRAM_DIR / "02-co-cau-sale-theo-mien.png",
        DIAGRAM_DIR / "03-bang-tong-ket-nhan-su.png",
        DIAGRAM_DIR / "04-ban-do-luong-du-lieu.png",
        DIAGRAM_DIR / "05-luong-1-marketing-tao-lead.png",
        DIAGRAM_DIR / "06-luong-2-crm-sale-nurture.png",
        DIAGRAM_DIR / "07-ho-tro-case-kho-sale-ctv.png",
        DIAGRAM_DIR / "08-luong-3-thanh-toan.png",
        DIAGRAM_DIR / "09-luong-4-onboarding.png",
        DIAGRAM_DIR / "10-luong-5-hoc-tap.png",
        DIAGRAM_DIR / "11-luong-6-cham-soc-hoc-vien.png",
        DIAGRAM_DIR / "12-luong-7-quan-ly-giang-vien.png",
        DIAGRAM_DIR / "13-luong-8-quan-ly-ctv.png",
        DIAGRAM_DIR / "14-luong-9-ke-toan.png",
        DIAGRAM_DIR / "15-phu-thuoc-cheo-cac-luong.png",
    ]
    render_org(1, paths[0])
    render_sales_structure(2, paths[1])
    render_headcount(3, paths[2])
    render_system_map(4, paths[3])
    render_vertical_flow(
        5,
        paths[4],
        DIAGRAM_TITLES[4],
        [
            "Quảng cáo Facebook / TikTok / Google",
            "Landing page: hsavnu.edu.vn",
            "Form đăng ký / form tư vấn → web database",
            "Sale nhập lead vào EZSale CRM",
        ],
        notes=[None, None, "thủ công, có độ trễ"],
        highlights={3: ORANGE},
        height=1350,
    )
    render_crm_nurture(6, paths[5])
    render_vertical_flow(
        7,
        paths[6],
        DIAGRAM_TITLES[6],
        [
            "Sale/CTV tư vấn học viên",
            "Phát sinh case khó / tư vấn sai / cần hỗ trợ chốt",
            "Quản lý hoặc nhân sự kinh nghiệm mở CRM kiểm tra thủ công",
            "Review ghi chú, trạng thái lead, lịch sử xử lý còn lại trong CRM",
            "Trao đổi riêng với Sale/CTV để hướng dẫn cách xử lý",
            "Case tương tự xuất hiện → quản lý phải review lại gần như từ đầu",
        ],
        notes=[None, "thủ công", "thủ công", None, "lặp lại"],
        highlights={1: ORANGE, 5: RED},
        height=1820,
    )
    render_vertical_flow(
        8,
        paths[7],
        DIAGRAM_TITLES[7],
        [
            "Học sinh chọn khóa trên web portal",
            "Thanh toán: chuyển khoản hoặc cổng thanh toán",
            "SePay webhook → web backend: payment_success",
            "Web database ghi nhận đơn hàng thanh toán",
            "Chuyển sang Luồng 4 - onboarding thủ công",
        ],
        notes=[None, "tự động", "tự động", "sau đây thủ công"],
        highlights={2: GREEN, 4: ORANGE},
        height=1500,
    )
    render_vertical_flow(
        9,
        paths[8],
        DIAGRAM_TITLES[8],
        [
            "Web đã ghi nhận thanh toán",
            "Nhân sự tạo SBD và ghi vào Google Sheet",
            "Gửi SBD + link nhóm Zalo lớp qua Zalo OA (gửi tay)",
            '"Duyệt học sinh": 1 chuyên viên kiểm tra Sheet có SBD + mã đơn hàng',
            "Add học sinh vào nhóm Zalo lớp; QLL gửi hướng dẫn học, link ClassIn/Zoom và tài liệu",
            "Học sinh tự cài/đăng nhập; nếu vướng thì hỏi trong nhóm Zalo và QLL hỗ trợ từng người",
        ],
        notes=["toàn bộ sau đây thủ công", "thủ công", "gửi tay", "single point of failure", "thủ công"],
        highlights={1: ORANGE, 2: ORANGE, 3: RED, 4: ORANGE, 5: ORANGE},
        height=1960,
    )
    render_vertical_flow(
        10,
        paths[9],
        DIAGRAM_TITLES[9],
        [
            "QLL/GV tạo hoặc vận hành lớp trên ClassIn",
            "QLL gửi hướng dẫn/link lớp trong nhóm Zalo lớp",
            "GV dạy trên ClassIn; một số lớp vẫn dùng Zoom chuyển tiếp hoặc dự phòng",
            "Điểm danh/dữ liệu học tập có trong ClassIn nhưng chưa sync tự động về hệ thống quản trị",
            "Học sinh hỏi đáp trong nhóm Zalo lớp sau giờ học",
            "QLL kiểm tra tình trạng học thủ công, đối chiếu ClassIn/Zalo/Sheet nếu cần",
        ],
        notes=[None, None, "chưa sync", None, "thủ công"],
        highlights={3: RED, 5: ORANGE},
        height=1880,
    )
    render_vertical_flow(
        11,
        paths[10],
        DIAGRAM_TITLES[10],
        [
            "Học sinh hỏi trong nhóm Zalo lớp",
            "GV / CTV trợ giảng / QLL phản hồi trực tiếp trong nhóm",
            "Nếu QLL phát hiện học sinh vắng nhiều buổi",
            "QLL nhắc thủ công qua Zalo cá nhân hoặc gọi điện",
            "Trước kỳ thi gửi tip thủ công; sau kỳ thi chưa có quy trình NPS hệ thống",
        ],
        notes=[None, "phụ thuộc phát hiện của QLL", "thủ công", "chưa chuẩn hóa NPS"],
        highlights={3: ORANGE, 4: RED},
        height=1560,
    )
    render_vertical_flow(
        12,
        paths[11],
        DIAGRAM_TITLES[11],
        [
            "Đầu mỗi đợt khai giảng: QLL Lead lập lịch dạy và gửi GV qua Zalo / Google Sheet",
            "GV nhận lịch, xác nhận",
            "Trong khóa: GV vào ClassIn dạy; Zoom là phương án chuyển tiếp/dự phòng",
            "Cuối tháng: GV hoặc QLL tổng hợp giờ dạy thủ công vào Sheet",
            "Kế toán chi đối chiếu Sheet và xử lý thù lao",
        ],
        notes=["Zalo/Sheet", None, "thủ công", "đối chiếu thủ công"],
        highlights={0: ORANGE, 3: RED, 4: ORANGE},
        height=1580,
    )
    render_vertical_flow(
        13,
        paths[12],
        DIAGRAM_TITLES[12],
        [
            "CTV giới thiệu học sinh và nhắn vào nhóm Zalo CTV",
            "Quản lý CTV ghi nhận vào Google Sheet",
            "Học sinh thanh toán; CTV hoặc quản lý CTV đối chiếu lại tay",
            "Cuối tháng: quản lý CTV tổng hợp số học sinh theo từng CTV",
            "Tính hoa hồng thủ công và chuyển kế toán chi",
            "Kế toán chi xử lý mạng lưới Sale/CTV khoảng 132-137 người",
        ],
        notes=["Zalo", "Sheet", "đối chiếu tay", "thủ công", "khối lượng lớn"],
        highlights={1: ORANGE, 2: ORANGE, 4: RED, 5: RED},
        height=1880,
    )
    render_accounting(14, paths[13])
    render_dependencies(15, paths[14])
    return paths


def set_style_font(doc: Document, style_name: str, font_name: str = "Arial", size: float | None = None, bold: bool | None = None, color: tuple[int, int, int] | None = None) -> None:
    style = doc.styles[style_name]
    style_font = style.font
    style_font.name = font_name
    if size is not None:
        style_font.size = Pt(size)
    if bold is not None:
        style_font.bold = bold
    if color is not None:
        style_font.color.rgb = RGBColor(*color)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)


def set_run_font(run, font_name: str = "Arial") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)


inline_pat = re.compile(r"(\[[^\]]+\]\([^)]+\)|`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
link_pat = re.compile(r"^\[([^\]]+)\]\(([^)]+)\)$")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    rfonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), "Arial")
    rpr.append(rfonts)
    new_run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, source: str) -> None:
    pos = 0
    for match in inline_pat.finditer(source):
        if match.start() > pos:
            run = paragraph.add_run(source[pos : match.start()])
            set_run_font(run)
        token = match.group(0)
        link_match = link_pat.match(token)
        if link_match:
            add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
        elif token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run)
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, "Consolas")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(80, 80, 80)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run)
            run.italic = True
        pos = match.end()
    if pos < len(source):
        run = paragraph.add_run(source[pos:])
        set_run_font(run)


def add_paragraph(doc: Document, text: str, style: str | None = None, align=None, space_after: float = 5):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    add_inline(p, text)
    return p


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def parse_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip().replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n") for cell in stripped.split("|")]


def is_separator_row(cells: Iterable[str]) -> bool:
    cells = list(cells)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 72, start: int = 72, bottom: int = 72, end: int = 72) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, pct: int = 5000) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(pct))
    tbl_w.set(qn("w:type"), "pct")


def add_table(doc: Document, raw_lines: list[str]) -> None:
    rows = [parse_table_row(line) for line in raw_lines]
    has_header = len(rows) > 1 and is_separator_row(rows[1])
    if has_header:
        header = rows[0]
        body = rows[2:]
    else:
        header = None
        body = rows
    all_rows = ([header] if header else []) + body
    if not all_rows:
        return
    col_count = max(len(row) for row in all_rows)
    table = doc.add_table(rows=len(all_rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_width(table)
    for r_idx, row in enumerate(all_rows):
        for c_idx in range(col_count):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            if r_idx == 0 and header is not None:
                set_cell_shading(cell, "1F4E79")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                add_inline(p, cell_text)
                for run in p.runs:
                    set_run_font(run)
                    run.font.size = Pt(8.2)
                    if r_idx == 0 and header is not None:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def syntax_starts(line: str) -> bool:
    stripped = line.strip()
    return (
        not stripped
        or stripped.startswith("#")
        or stripped.startswith("```")
        or stripped in {"---", "***", "___"}
        or stripped.startswith(">")
        or is_table_line(stripped)
        or re.match(r"^\s*[-*+]\s+", line)
        or re.match(r"^\s*\d+[.)]\s+", line)
        or re.match(r"^\s*\*\*[^*]+:\*\*", line)
    )


def add_diagram(doc: Document, image_path: Path, title: str, index: int) -> None:
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    run = caption.add_run(f"Hình {index}. {title}")
    set_run_font(run)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*NAVY)
    picture_para = doc.add_paragraph()
    picture_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_para.paragraph_format.space_after = Pt(8)
    run = picture_para.add_run()
    with Image.open(image_path) as img:
        ratio = img.height / img.width
    max_width_cm = 25.5
    max_height_cm = 17.0
    width_cm = min(max_width_cm, max_height_cm / ratio)
    run.add_picture(str(image_path), width=Cm(width_cm))


def build_doc(diagram_paths: list[Path]) -> None:
    text = SRC.read_text(encoding="utf-8-sig")
    doc = Document()
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)
    props = doc.core_properties
    props.title = "HSA Education - As-Is Operations Analysis"
    props.subject = "Q2/2026 as-is operations analysis with redrawn model PNGs"
    props.author = "HSA Education"
    set_style_font(doc, "Normal", size=10.5)
    for name, size, color in [
        ("Title", 20, NAVY),
        ("Heading 1", 18, NAVY),
        ("Heading 2", 14, NAVY),
        ("Heading 3", 12, (45, 45, 45)),
        ("Heading 4", 11, (45, 45, 45)),
        ("Intense Quote", 10, (90, 90, 90)),
    ]:
        if name in doc.styles:
            set_style_font(doc, name, size=size, bold=True if "Heading" in name or name == "Title" else None, color=color)
    for style_name in ["List Bullet", "List Bullet 2", "List Number", "List Number 2"]:
        if style_name in doc.styles:
            set_style_font(doc, style_name, size=10.5)
    lines = text.splitlines()
    i = 0
    diagram_index = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("```"):
            while i + 1 < len(lines) and not lines[i + 1].strip().startswith("```"):
                i += 1
            i += 2
            diagram_index += 1
            if diagram_index <= len(diagram_paths):
                add_diagram(doc, diagram_paths[diagram_index - 1], DIAGRAM_TITLES[diagram_index - 1], diagram_index)
            continue
        if stripped in {"---", "***", "___"}:
            add_horizontal_rule(doc)
            i += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            title = heading.group(2).strip()
            if level == 1:
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(4)
                add_inline(p, title)
            else:
                p = doc.add_heading("", level=min(level, 4))
                p.paragraph_format.space_before = Pt(8 if level <= 2 else 5)
                p.paragraph_format.space_after = Pt(4)
                add_inline(p, title)
            i += 1
            continue
        if is_table_line(stripped):
            table_lines: list[str] = []
            while i < len(lines) and is_table_line(lines[i].strip()):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue
        if stripped.startswith(">"):
            quote_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            p = add_paragraph(doc, " ".join(quote_lines), style="Intense Quote", space_after=6)
            p.paragraph_format.left_indent = Cm(0.35)
            continue
        bullet = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        if bullet:
            indent = len(bullet.group(1).replace("\t", "    "))
            style = "List Bullet 2" if indent >= 2 and "List Bullet 2" in doc.styles else "List Bullet"
            add_paragraph(doc, bullet.group(2).strip(), style=style, space_after=2)
            i += 1
            continue
        number = re.match(r"^(\s*)\d+[.)]\s+(.*)$", line)
        if number:
            indent = len(number.group(1).replace("\t", "    "))
            style = "List Number 2" if indent >= 2 and "List Number 2" in doc.styles else "List Number"
            add_paragraph(doc, number.group(2).strip(), style=style, space_after=2)
            i += 1
            continue
        if re.match(r"^\s*\*\*[^*]+:\*\*", line):
            add_paragraph(doc, stripped, space_after=1)
            i += 1
            continue
        para_lines = [stripped]
        i += 1
        while i < len(lines) and not syntax_starts(lines[i]):
            para_lines.append(lines[i].strip())
            i += 1
        add_paragraph(doc, " ".join(para_lines), space_after=5)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Trang ")
    set_run_font(run)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instr)
    run._r.append(field_end)
    doc.save(OUT)


def main() -> None:
    diagram_paths = generate_diagrams()
    build_doc(diagram_paths)
    print(f"created_docx={OUT}")
    print(f"created_pngs={len(diagram_paths)}")
    for item in diagram_paths:
        print(item)


if __name__ == "__main__":
    main()
